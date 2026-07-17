import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import comtypes.client

def set_cell_background(cell, fill_hex):
    """Set the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def generate_report():
    doc = Document()
    
    # Configure page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Color Palette (Indigo/Slate scheme)
    COLOR_PRIMARY = RGBColor(79, 70, 229)   # Indigo #4f46e5
    COLOR_SECONDARY = RGBColor(15, 23, 42)  # Slate-900 #0f172a
    COLOR_TEXT = RGBColor(51, 65, 85)       # Slate-700 #334155
    COLOR_MUTED = RGBColor(100, 116, 139)   # Slate-500 #64748b

    # Style Helpers
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run(text)
        run.font.name = 'Trebuchet MS'
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(36)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_MUTED
        return p

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Trebuchet MS'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Trebuchet MS'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_body_text(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXT
        return p

    def add_bullet_item(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXT
        return p

    def add_code_block(filepath):
        # Fetch file content
        if not os.path.exists(filepath):
            return
        
        with open(filepath, "r", encoding="utf-8") as f:
            code = f.read()

        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Inches(6.5)
        
        cell = table.cell(0, 0)
        set_cell_background(cell, "F1F5F9")  # Slate-100 background
        
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.right_indent = Inches(0.1)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        run = p.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.0)
        run.font.color.rgb = RGBColor(15, 23, 42) # Slate-900 code text

    def add_image_block(image_path, caption):
        if not os.path.exists(image_path):
            return
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(12)
        p_img.paragraph_format.space_after = Pt(4)
        r_img = p_img.add_run()
        r_img.add_picture(image_path, width=Inches(6.0))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        r_cap = p_cap.add_run("Figure: " + caption)
        r_cap.font.name = 'Calibri'
        r_cap.font.size = Pt(9.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = COLOR_MUTED

    # --- Title Page Content ---
    add_title("✦ Syntactic Intelligence AI Assistant")
    add_subtitle(
        "Complete Technical Report & Project Documentation\n"
        "GitHub Repository: https://github.com/abhinavsai2006/Syntactic-Intelligence.git\n"
        "Major Project: Prompt Engineering | Powered by NVIDIA NIM Llama-3.1-8B-Instruct API"
    )

    # --- Section 1: Introduction & Goals ---
    add_heading_1("1. Project Goals & Learning Objectives")
    add_body_text(
        "The primary goal of the Syntactic Intelligence project is to design and implement a web-based AI assistant "
        "incorporating prompt design variance, output evaluation, and user feedback logs. This project aims to enhance "
        "students' capabilities to formulate systemic prompts that inspire structured thinking, logical evaluation, and "
        "entrepreneurial creativity."
    )
    add_body_text("The assistant is configured to perform three distinct tasks based on Llama-3.1-8B-Instruct LLM responses:")
    add_bullet_item("Answer Questions: Factual Q&A distilled into concise summaries or detailed explanations.")
    add_bullet_item("Summarize Text: Distill long articles or documents while preserving nuance.")
    add_bullet_item("Generate Creative Content: Brainstorm and write stories or poems based on specific input parameters.")

    # --- Section 2: Technical Stack & Architecture ---
    add_heading_1("2. Technical Stack & Component Design")
    add_body_text(
        "The project is structured under a flat, light-corporate design metaphor (without glassmorphism) using "
        "Tailwind CSS for maximum structural density and clean border separations."
    )
    
    add_heading_2("2.1 Frontend Architecture")
    add_bullet_item("Tailwind CSS CDN Integration: Used to configure responsive columns, utility inputs, and flat border cards.")
    add_bullet_item("Asynchronous AJAX Client: Utilizes standard browser fetch() calls to submit forms and rate feedbacks without page reloading.")
    add_bullet_item("Light Slate Canvas: Background styled in `#f7f9fb` with `#ffffff` container cards to ensure a sleek developer interface.")
    
    add_heading_2("2.2 Backend Architecture")
    add_bullet_item("Flask Framework: Coordinates routing pipelines between the HTML workspace, API calling helpers, and CSV log writers.")
    add_bullet_item("NVIDIA NIM API Integration: Interfaces with the meta/llama-3.1-8b-instruct chat completion endpoint.")
    add_bullet_item("Locking Logging System: Utilizes threading.Lock() primitives to write logs safely to feedback_log.csv.")

    add_heading_2("2.3 System Pipeline Flowchart")
    add_body_text(
        "The diagram below details the end-to-end data lifecycle, from initial UI user selection inputs, through Flask server "
        "route handler transformations and NVIDIA NIM Llama-3.1 LLM request formatting, to the thread-locked CSV metrics logger:"
    )
    add_image_block("flowchart.png", "System Pipeline Flowchart detailing the request-response lifecycle.")

    # --- Section 3: App Interfaces & Screenshots ---
    add_heading_1("3. Application Working & Interface Mockups")
    add_body_text("Here are the live screenshots capturing the assistant's workflow and response logs:")
    
    add_heading_2("3.1 Q&A Response Interface")
    add_body_text("The screen below showcases a user asking 'what is ai' under the Q&A layout. The prompt variation chosen is 'Bullet-Point Facts':")
    add_image_block("qa_screen.png", "Llama-3.1-8b NIM rendering structured markdown bullet-points in the flat workspace.")

    add_heading_2("3.2 Feedback Summary Dashboard")
    add_body_text("The screen below displays the aggregates page (/feedback-summary) plotting total interactions, helpful metrics, and rating distributions:")
    add_image_block("analytics_screen.png", "Analytics Dashboard displaying aggregate satisfaction rates and per-prompt breakdowns.")

    # --- Section 4: Source Code Documentation ---
    add_heading_1("4. Complete Source Code & Implementation Details")
    
    # 4.1 app.py
    add_heading_2("4.1 Flask Server Routing (app.py)")
    add_body_text(
        "File Description:\n"
        "The app.py file serves as the main application controller. It imports requests to handle HTTPS payloads, csv to log ratings, "
        "and threading to lock the logger file during multi-user write scenarios. It declares the global PROMPTS, SYSTEM_PROMPTS, "
        "and labels, and connects user requests with the NVIDIA NIM chat completions model (meta/llama-3.1-8b-instruct) endpoints."
    )
    add_body_text("Code Segment:")
    add_code_block("app.py")
    add_body_text(
        "Execution Details:\n"
        "Upon receipt of client POST requests on /generate, the helper function call_nvidia formats the template with user inputs, "
        "queries Llama-3.1-8b, and returns structured markdown payloads. Thumbs up/down feedback ratings call /feedback, appending "
        "them safely to the local CSV logger."
    )

    # 4.2 index.html
    add_heading_2("4.2 Main HTML Workspace (templates/index.html)")
    add_body_text(
        "File Description:\n"
        "The index.html file provides the core workspace layout, styled in a flat light-corporate style using Tailwind CSS. "
        "It sets up tabs for selecting functions, displays prompt style selectors, includes a character-counted input textarea, "
        "and provides copying mechanisms and thumbs-up/down ratings for user response cards."
    )
    add_body_text("Code Segment:")
    add_code_block("templates/index.html")
    add_body_text(
        "Execution Details:\n"
        "All forms use event.preventDefault() and fetch() calls. When a user changes a tab, selectFunction() updates "
        "dropdown options and descriptions. Responses are rendered on the page using a lightweight custom markdown parsing helper."
    )

    # 4.3 feedback_summary.html
    add_heading_2("4.3 Analytics HTML template (templates/feedback_summary.html)")
    add_body_text(
        "File Description:\n"
        "The feedback_summary.html file provides the analytics dashboard layout. It displays cards for total interactions, "
        "helpful scores, dissatisfaction totals, satisfaction rates, and renders progress bars and per-style breakdown tables."
    )
    add_body_text("Code Segment:")
    add_code_block("templates/feedback_summary.html")
    add_body_text(
        "Execution Details:\n"
        "The file utilizes Jinja2 syntax to iterate over the statistics aggregates calculated by the app.py helper get_feedback_stats(). "
        "It provides a clear overview of prompt helpfulness, letting users refine system prompts based on empirical metrics."
    )

    # 4.4 requirements.txt
    add_heading_2("4.4 Setup Requirements (requirements.txt)")
    add_body_text(
        "File Description:\n"
        "Declares python libraries necessary to run the web application server locally. Uses flask for routing, "
        "requests to call the NIM endpoint, and python-dotenv to isolate api keys from source control repositories."
    )
    add_body_text("Code Segment:")
    add_code_block("requirements.txt")

    # 4.5 .gitignore
    add_heading_2("4.5 Git Exclusions (.gitignore)")
    add_body_text(
        "File Description:\n"
        "Excludes python cache objects, virtual environments (.venv), user feedback logs, and sensitive credential variables "
        "(.env) from remote repositories."
    )
    add_body_text("Code Segment:")
    add_code_block(".gitignore")

    # Save Document
    docx_path = "AI_Assistant_Technical_Report.docx"
    pdf_path = "AI_Assistant_Technical_Report.pdf"
    
    doc.save(docx_path)
    print(f"Word document saved to: {docx_path}")

    # Compile to PDF using Word Application via COM
    try:
        print("Compiling to PDF...")
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc_obj = word.Documents.Open(os.path.abspath(docx_path))
        doc_obj.SaveAs(os.path.abspath(pdf_path), FileFormat=17) # 17 is for PDF
        doc_obj.Close()
        word.Quit()
        print(f"PDF document saved to: {pdf_path}")
    except Exception as e:
        print(f"Failed to compile to PDF: {str(e)}")

if __name__ == "__main__":
    generate_report()
